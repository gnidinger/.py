import pandas as pd
import random
import docx
from docx.shared import Pt


def excel_to_doc_v4(file_name, num_rows, input_columns, output_column, doc_name):
    df = pd.read_excel(file_name)

    df = df[df.iloc[:, 7].notna()]

    indices = select_rows(df, num_rows)

    # Select rows based on the indices
    selected_rows_df = df.iloc[indices]

    doc = docx.Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    doc_content = ""

    for index, row in selected_rows_df.iterrows():
        input_text = ", ".join(str(row.iloc[i]) for i in input_columns)
        output_text = str(row.iloc[output_column])
        # paragraph = doc.add_paragraph()

        doc_content += "\\n###\\n제품 설명: " + input_text + "\\n" + "출력: " + output_text

    with open(doc_name, "w", encoding="utf-8-sig") as f:
        f.write(doc_content)

        # doc.add_paragraph("\\n###\\n" + input_text + "\\n" + "출력: " + output_text)

    # doc.save(doc_name)


def select_rows(df, num_rows):
    total_rows = len(df)

    if num_rows > total_rows:
        raise ValueError(
            "num_rows cannot be greater than the total number of rows in the DataFrame"
        )

    selected_indices = random.sample(range(total_rows), num_rows)

    return selected_indices


# specify columns by their indices, starting from 0
input_columns = [22]  # replace with your input column indices
output_column = 23  # replace with your output column index


def excel_to_doc_v5(file_path, num_rows, input_columns, output_column):
    # Load the Excel file into a pandas DataFrame
    df = pd.read_excel(file_path)

    # Randomly select num_rows rows
    selected_rows = df.sample(n=num_rows)

    # Open the output file
    with open("/Users/gnidinger/Downloads/HSAD_DataSet/LG1.txt", "w", encoding="utf-8") as f:
        # Iterate over the selected rows
        for _, row in selected_rows.iterrows():
            # Generate the input text
            input_text = ", ".join(f"{df.columns[i]}: {row.iloc[i]}" for i in input_columns)
            # Write the input text to the file
            f.write(f"\n{input_text}\n")

            # Get the output text
            output_text = row.iloc[output_column]
            # Write the output text to the file
            f.write(f"출력: {output_text}\n")

            # Write a newline to the file
            f.write("\n")


# excel_to_doc_v5(
#     "/Users/gnidinger/Downloads/HSAD_DataSet/LG1.xlsx",
#     11,
#     [7, 8, 19, 9, 10, 11, 21, 20],
#     22,
# )


excel_to_doc_v4(
    "/Users/gnidinger/Documents/Python/HSAD_DataSet/LG1.xlsx",
    30,
    [7, 8, 19, 9, 10, 11, 21, 20],
    2,
    "/Users/gnidinger/Documents/Python/HSAD_DataSet/LG1.doc",
)
